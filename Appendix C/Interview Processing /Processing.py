import re
import docx
import argparse
import glob
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.tag import pos_tag
from nltk.stem import WordNetLemmatizer

# Download necessary NLTK resources
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('taggers/averaged_perceptron_tagger')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('corpora/wordnet')
except LookupError:
    print("Downloading necessary NLTK resources...")
    nltk.download('punkt')
    nltk.download('averaged_perceptron_tagger')
    nltk.download('stopwords')
    nltk.download('wordnet')

def clean_transcript(input_file, output_file=None):
    """
    Clean a transcript from a Word file by:
    - Removing timestamps
    - Organizing content into logical sections
    - Cleaning up speech disfluencies
    - Improving conversation flow
    - Clearly identifying speakers
    - Preserving technical terminology
    - Clarifying ambiguous statements
    
    Args:
        input_file (str): Path to the input Word file
        output_file (str, optional): Path to save the cleaned transcript. 
                                    If None, will use input_file_cleaned.docx
    
    Returns:
        str: Path to the saved cleaned transcript
    """
    # Set default output file if not provided
    if output_file is None:
        file_name, file_ext = os.path.splitext(input_file)
        output_file = f"{file_name}_cleaned.docx"
    
    print(f"Reading transcript from: {input_file}")
    
    try:
        # Load the document
        doc = Document(input_file)
        full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract the cleaned content
        cleaned_text = process_transcript(full_text)
        
        # Create a new document with the cleaned content
        create_formatted_document(cleaned_text, output_file)
        
        print(f"Cleaned transcript saved to: {output_file}")
        return output_file
        
    except Exception as e:
        print(f"Error processing transcript: {str(e)}")
        return None

def process_transcript(text):
    """Process the transcript text to clean and organize it"""
    
    # Step 1: Remove timestamps (assuming formats like [00:15] or (12:45) or 10:30 or standalone 00:00:05)
    text = re.sub(r'[\[\(]?\d{1,2}:\d{2}(:\d{2})?[\]\)]?', '', text)
    
    # Step 2: Identify and standardize speaker labels
    # First, handle Speaker notation with asterisks pattern like "*Speaker 1: *" or "*Speaker 2: *"
    text = re.sub(r'\*Speaker 1:\s*\*', 'Interviewer: ', text, flags=re.IGNORECASE)
    text = re.sub(r'\*Speaker 2:\s*\*', 'Expert: ', text, flags=re.IGNORECASE)
    
    # Then handle other common speaker notations
    text = re.sub(r'\b(?:Interviewer|Questioner|Q|Interviewier)[\s:]+', 'Interviewer: ', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(?:Michael|Expert|Subject|Respondent|A|Mike)[\s:]+', 'Expert: ', text, flags=re.IGNORECASE)
    
    # Step 3: Clean up speech disfluencies
    # Remove common filler words and phrases
    disfluencies = [r'\bum+\b', r'\buh+\b', r'\blike\b', r'\byou know\b', 
                   r'\bI mean\b', r'\bso\b', r'\bjust\b', r'\bkind of\b', 
                   r'\bsort of\b', r'\bliterally\b', r'\bbasically\b']
    
    for pattern in disfluencies:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # Remove repeated words (e.g., "the the", "I I", etc.)
    text = re.sub(r'\b(\w+)(\s+\1\b)+', r'\1', text, flags=re.IGNORECASE)
    
    # Fix broken sentences with periods followed by lowercase letters
    text = re.sub(r'\.(\s+[a-z])', lambda m: '. ' + m.group(1).strip().capitalize(), text)
    
    # Remove any remaining asterisks
    text = re.sub(r'\*', '', text)
    
    # Step 4: Fix incomplete sentences and improve flow
    # This is more complex and might require manual review
    
    # Step 5: Split the text into sections based on topic changes
    # This is more advanced and may require content understanding
    
    # For now, let's split the text into paragraphs by speaker
    paragraphs = []
    current_speaker = None
    current_paragraph = []
    
    # Split by lines first
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this line starts with a speaker label
        speaker_match = re.match(r'^(Interviewer|Michael):\s*(.*)', line, re.IGNORECASE)
        
        if speaker_match:
            # If we have a previous paragraph, add it
            if current_paragraph:
                paragraphs.append((current_speaker, ' '.join(current_paragraph)))
                current_paragraph = []
            
            # Start a new paragraph with the new speaker
            current_speaker = speaker_match.group(1)
            content = speaker_match.group(2).strip()
            if content:
                current_paragraph.append(content)
        else:
            # Continue with the current paragraph
            if current_speaker:
                current_paragraph.append(line)
            else:
                # If no speaker identified yet, assume it's part of the introduction or metadata
                paragraphs.append((None, line))
    
    # Add the last paragraph if there is one
    if current_paragraph:
        paragraphs.append((current_speaker, ' '.join(current_paragraph)))
    
    # Combine consecutive paragraphs from the same speaker
    combined_paragraphs = []
    current_speaker = None
    current_content = []
    
    for speaker, content in paragraphs:
        if speaker == current_speaker:
            current_content.append(content)
        else:
            if current_speaker is not None:
                combined_paragraphs.append((current_speaker, ' '.join(current_content)))
            current_speaker = speaker
            current_content = [content]
    
    # Add the last group
    if current_content:
        combined_paragraphs.append((current_speaker, ' '.join(current_content)))
    
    # Clean up each paragraph's content
    cleaned_paragraphs = []
    for speaker, content in combined_paragraphs:
        # Remove multiple spaces
        content = re.sub(r'\s+', ' ', content).strip()
        
        # Fix punctuation spacing
        content = re.sub(r'\s*([,.;:!?])', r'\1', content)
        
        # Make sure sentences end with proper punctuation
        if content and content[-1] not in ".!?":
            content += "."
            
        # Capitalize the first letter of each sentence
        content = '. '.join(s.strip().capitalize() for s in content.split('. '))
        
        cleaned_paragraphs.append((speaker, content))
    
    # Create sections based on topic changes
    # This is a simple approach - a more advanced one would involve NLP
    section_break_phrases = [
        "let's move on to", "next topic", "another question",
        "next question", "changing subjects", "moving forward"
    ]
    
    # Initialize sections
    sections = []
    current_section = []
    
    for speaker, content in cleaned_paragraphs:
        # Check if this paragraph indicates a topic change
        is_section_break = False
        lower_content = content.lower()
        
        for phrase in section_break_phrases:
            if phrase in lower_content:
                is_section_break = True
                # Split at the section break phrase
                break
        
        if is_section_break and current_section:
            sections.append(current_section)
            current_section = []
        
        current_section.append((speaker, content))
    
    # Add the last section
    if current_section:
        sections.append(current_section)
    
    # Format the final transcript
    formatted_transcript = []
    
    for i, section in enumerate(sections):
        # Add a section header if this isn't the introduction
        if i > 0:
            formatted_transcript.append(f"Section {i}: Topic {i}")
        
        # Add the paragraphs in this section
        for speaker, content in section:
            if speaker:
                formatted_transcript.append(f"{speaker}: {content}")
            else:
                formatted_transcript.append(content)
        
        # Add a separator between sections
        if i < len(sections) - 1:
            formatted_transcript.append("")
    
    return "\n".join(formatted_transcript)

def create_formatted_document(text, output_file):
    """Create a nicely formatted Word document with the cleaned transcript"""
    doc = Document()
    
    # Set up document styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.size = Pt(16)
    title_font.bold = True
    title_paragraph_format = title_style.paragraph_format
    title_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_paragraph_format.space_after = Pt(12)
    
    # Section style
    section_style = styles.add_style('Section', WD_STYLE_TYPE.PARAGRAPH)
    section_font = section_style.font
    section_font.size = Pt(14)
    section_font.bold = True
    section_paragraph_format = section_style.paragraph_format
    section_paragraph_format.space_before = Pt(12)
    section_paragraph_format.space_after = Pt(6)
    
    # Speaker style
    speaker_style = styles.add_style('Speaker', WD_STYLE_TYPE.CHARACTER)
    speaker_font = speaker_style.font
    speaker_font.bold = True
    
    # Add title
    title = doc.add_paragraph("Cleaned Interview Transcript", style='Title')
    
    # Process the text line by line
    lines = text.split('\n')
    
    for line in lines:
        if not line.strip():
            # Add an empty paragraph for spacing
            doc.add_paragraph()
        elif line.startswith('Section'):
            # This is a section header
            doc.add_paragraph(line, style='Section')
        else:
            # Check if this is a speaker line
            speaker_match = re.match(r'^(Interviewer|Expert):\s*(.*)', line)
            
            if speaker_match:
                speaker = speaker_match.group(1)
                content = speaker_match.group(2)
                
                p = doc.add_paragraph()
                p.add_run(f"{speaker}: ").bold = True
                p.add_run(content)
            else:
                # Regular paragraph
                doc.add_paragraph(line)
    
    # Save the document
    doc.save(output_file)

def main():
    parser = argparse.ArgumentParser(description='Clean a transcript from a Word file.')
    parser.add_argument('input_file', help='Path to the input Word file')
    parser.add_argument('--output', '-o', help='Path to save the cleaned transcript')
    parser.add_argument('--example', '-e', action='store_true', help='Process the example text instead of a file')
    
    args = parser.parse_args()
    
    if args.example:
        # Example text processing
        example_text = """*Speaker 1: *So can you please introduce yourself a little bit?
00:00:05 *Speaker 2: *. Uh, I'm an associate professor here. I'd like to you. I've been here since, uh, 2009, actually. So I was a very long time. And, um, my research area is, uh, uh, software verification. Uh, and I mostly work with theoretical in, uh, in the theoretical computer science area. I'm also a member of the CSat center, which is our center for information. Uh, trust. Uh, no. Sorry. Information security and trust. And, um, I am, uh, also had in the, uh, our master's in computer science, as I do."""
        
        result = process_example_transcript(example_text)
        print(result)
    else:
        # Process from file
        clean_transcript(args.input_file, args.output)

if __name__ == "__main__":
    main()

if __name__ == "__main__":
    main()